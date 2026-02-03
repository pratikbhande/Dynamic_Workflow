"""Report Generator Tool - Create professional DOCX/PDF reports"""
from typing import Dict, Any, List, Optional
from ..base_tool import (
    BasePredefinedTool,
    ToolMetadata,
    CredentialRequirement,
    InputParameter,
    OutputSchema,
    ToolExecutionResult,
    ToolCategory
)


class ReportGeneratorTool(BasePredefinedTool):
    """
    Report Generator Tool - Create downloadable reports
    
    Features:
    - DOCX and PDF generation
    - Embedded charts and visualizations
    - Professional formatting
    - Table generation
    - Custom styling
    """
    
    def get_metadata(self) -> ToolMetadata:
        return ToolMetadata(
            name="report_generator",
            display_name="Professional Report Generator",
            description="Generate downloadable DOCX or PDF reports with charts and professional formatting",
            category=ToolCategory.REPORT,
            tags=["report", "docx", "pdf", "charts", "visualization"]
        )
    
    def get_required_credentials(self) -> List[CredentialRequirement]:
        return []  # No credentials required
    
    def get_input_parameters(self) -> List[InputParameter]:
        return [
            InputParameter(
                name="title",
                type="string",
                description="Report title",
                required=True
            ),
            InputParameter(
                name="sections",
                type="array",
                description="List of sections with heading and content",
                required=True
            ),
            InputParameter(
                name="data",
                type="object",
                description="Data for visualizations (optional)",
                required=False
            ),
            InputParameter(
                name="charts",
                type="array",
                description="Chart specifications (optional)",
                required=False
            ),
            InputParameter(
                name="output_format",
                type="string",
                description="Output format",
                required=False,
                default="docx",
                options=["docx", "pdf"]
            ),
            InputParameter(
                name="include_toc",
                type="boolean",
                description="Include table of contents",
                required=False,
                default=False
            )
        ]
    
    def get_output_schema(self) -> OutputSchema:
        return OutputSchema(
            type="object",
            description="Generated report details",
            properties={
                "file_path": "Path to generated report",
                "file_name": "Name of generated file",
                "format": "Report format (docx/pdf)",
                "sections": "Number of sections",
                "charts": "Number of charts included"
            }
        )
    
    async def execute(
        self,
        inputs: Dict[str, Any],
        credentials: Optional[Dict[str, str]] = None
    ) -> ToolExecutionResult:
        """Execute report generation"""
        
        try:
            title = inputs["title"]
            sections = inputs["sections"]
            data = inputs.get("data", {})
            charts = inputs.get("charts", [])
            output_format = inputs.get("output_format", "docx")
            include_toc = inputs.get("include_toc", False)
            
            print(f"\n📊 Report Generator Starting...")
            print(f"   Title: {title}")
            print(f"   Sections: {len(sections)}")
            print(f"   Charts: {len(charts)}")
            print(f"   Format: {output_format}")
            
            # Step 1: Generate charts if requested
            chart_paths = []
            if charts:
                chart_paths = await self._generate_charts(charts, data)
                print(f"   ✅ Generated {len(chart_paths)} charts")
            
            # Step 2: Generate report
            if output_format == "docx":
                file_path = await self._generate_docx(
                    title=title,
                    sections=sections,
                    chart_paths=chart_paths,
                    include_toc=include_toc
                )
            else:
                file_path = await self._generate_pdf(
                    title=title,
                    sections=sections,
                    chart_paths=chart_paths
                )
            
            print(f"   ✅ Report generated: {file_path}")
            
            # Copy to outputs directory for download
            import shutil
            import os
            output_dir = "/mnt/user-data/outputs"
            os.makedirs(output_dir, exist_ok=True)
            
            file_name = os.path.basename(file_path)
            output_path = os.path.join(output_dir, file_name)
            shutil.copy(file_path, output_path)
            
            return ToolExecutionResult(
                success=True,
                output={
                    "file_path": output_path,
                    "file_name": file_name,
                    "format": output_format,
                    "sections": len(sections),
                    "charts": len(chart_paths)
                },
                metadata={
                    "title": title,
                    "download_url": f"/outputs/{file_name}"
                }
            )
            
        except Exception as e:
            import traceback
            return ToolExecutionResult(
                success=False,
                output=None,
                error=f"Report generation failed: {str(e)}\n{traceback.format_exc()}"
            )
    
    async def _generate_charts(
        self,
        charts: List[Dict[str, Any]],
        data: Dict[str, Any]
    ) -> List[str]:
        """Generate chart images"""
        import matplotlib.pyplot as plt
        import matplotlib
        matplotlib.use('Agg')
        import uuid
        import os
        
        chart_paths = []
        chart_dir = "/app/data/uploads/charts"
        os.makedirs(chart_dir, exist_ok=True)
        
        for chart_spec in charts:
            try:
                chart_type = chart_spec.get("type", "bar")
                chart_data = chart_spec.get("data", data)
                chart_title = chart_spec.get("title", "Chart")
                
                # Create figure
                plt.figure(figsize=(10, 6))
                
                if chart_type == "bar":
                    plt.bar(chart_data.keys(), chart_data.values())
                elif chart_type == "line":
                    plt.plot(list(chart_data.keys()), list(chart_data.values()), marker='o')
                elif chart_type == "pie":
                    plt.pie(chart_data.values(), labels=chart_data.keys(), autopct='%1.1f%%')
                elif chart_type == "scatter":
                    plt.scatter(list(chart_data.keys()), list(chart_data.values()))
                
                plt.title(chart_title, fontsize=14, fontweight='bold')
                plt.tight_layout()
                
                # Save
                chart_path = os.path.join(chart_dir, f"chart_{uuid.uuid4().hex[:8]}.png")
                plt.savefig(chart_path, dpi=300, bbox_inches='tight')
                plt.close()
                
                chart_paths.append(chart_path)
                
            except Exception as e:
                print(f"   ⚠️  Chart generation failed: {e}")
        
        return chart_paths
    
    async def _generate_docx(
        self,
        title: str,
        sections: List[Dict[str, Any]],
        chart_paths: List[str],
        include_toc: bool
    ) -> str:
        """Generate DOCX report"""
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
        import uuid
        import os
        from datetime import datetime
        
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = title
        doc.core_properties.author = "AI Report Generator"
        
        # Title
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date
        date_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.runs[0].font.size = Pt(10)
        date_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        
        doc.add_page_break()
        
        # Table of contents placeholder
        if include_toc:
            doc.add_heading("Table of Contents", 1)
            doc.add_paragraph("(Table of contents would be here)")
            doc.add_page_break()
        
        # Add sections
        chart_index = 0
        for section in sections:
            heading = section.get("heading", "Section")
            content = section.get("content", "")
            
            # Section heading
            doc.add_heading(heading, 1)
            
            # Section content
            if content:
                # Split into paragraphs
                paragraphs = content.split('\n\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        para = doc.add_paragraph(para_text.strip())
                        para.style.font.size = Pt(11)
            
            # Add chart if available
            if chart_index < len(chart_paths):
                doc.add_paragraph()  # Spacing
                doc.add_picture(chart_paths[chart_index], width=Inches(6))
                chart_index += 1
            
            doc.add_paragraph()  # Spacing between sections
        
        # Save
        output_path = f"/app/data/uploads/report_{uuid.uuid4().hex[:8]}.docx"
        doc.save(output_path)
        
        return output_path
    
    async def _generate_pdf(
        self,
        title: str,
        sections: List[Dict[str, Any]],
        chart_paths: List[str]
    ) -> str:
        """Generate PDF report"""
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, PageBreak
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib.enums import TA_CENTER
        import uuid
        from datetime import datetime
        
        output_path = f"/app/data/uploads/report_{uuid.uuid4().hex[:8]}.pdf"
        
        doc = SimpleDocTemplate(output_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor='#1a1a1a',
            spaceAfter=30,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        story.append(Paragraph(title, title_style))
        
        # Date
        date_style = ParagraphStyle(
            'DateStyle',
            parent=styles['Normal'],
            fontSize=10,
            textColor='#808080',
            alignment=TA_CENTER
        )
        story.append(Paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}", date_style))
        story.append(Spacer(1, 0.5*inch))
        
        # Sections
        chart_index = 0
        for section in sections:
            heading = section.get("heading", "Section")
            content = section.get("content", "")
            
            # Section heading
            story.append(Paragraph(heading, styles['Heading2']))
            story.append(Spacer(1, 0.2*inch))
            
            # Section content
            if content:
                paragraphs = content.split('\n\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        story.append(Paragraph(para_text.strip(), styles['BodyText']))
                        story.append(Spacer(1, 0.1*inch))
            
            # Add chart if available
            if chart_index < len(chart_paths):
                story.append(Spacer(1, 0.2*inch))
                story.append(Image(chart_paths[chart_index], width=6*inch, height=4*inch))
                chart_index += 1
            
            story.append(Spacer(1, 0.3*inch))
        
        # Build PDF
        doc.build(story)
        
        return output_path